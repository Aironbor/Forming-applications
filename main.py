import logging
import os
import subprocess
import sys

import PyQt5
from PyQt5 import QtWidgets
from PyQt5.QtCore import Qt
from PyQt5 import uic
from PyQt5 import sip
from PyQt5 import QtCore, QtWidgets, QtGui
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QIcon, QPixmap
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

from images import images_store
from message_widgets import ErrorAddReport, MessageDialogWindow
import os
from docx.templates import *
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PyQt5 import QtWidgets, uic, QtCore
from PyQt5.QtWidgets import QFileDialog
import sys
from datetime import datetime
from PIL import Image
from PyPDF2 import PdfMerger
import docx2pdf

# Настройка логгирования
log_dir = "logging"
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

logging.basicConfig(filename=os.path.join(log_dir, 'logfile.txt'),
                    level=logging.INFO,
                    format='%(asctime)s:%(levelname)s:%(message)s')


class MainProgramMenu(QtWidgets.QMainWindow):
    def __init__(self, parent=None, flag=Qt.Window):
        super().__init__(parent, flag)
        uic.loadUi('ui/menu_applications.ui', self)
        self.dict_product_and_info = {}  # Инициализируем словарь с данными о добавленых продуктах
        self.files_glob = []
        self.product_index = 1
        self.load_interface()
        self.MainWindow = ''

    def load_interface(self) -> None:
        list_of_constr_solutions = ["Изделие", "Рамное", "Модуль (транспак)"]
        self.quant_sb.setMaximum(100000)
        self.cnstr_solut_combobox.addItems(list_of_constr_solutions)
        self.cnstr_solut_combobox.currentIndexChanged.connect(self.constr_solution_changed)
        self.is_it_standart_checkbox.stateChanged.connect(self.gabarits_checkbox_changed_state)
        self.is_drawings_checkbox.stateChanged.connect(self.have_progect)
        self.listWidget_products.currentItemChanged.connect(self.change_product_in_list)
        self.add_prod_btn.clicked.connect(self.add_product_to_listwidget)
        self.choise_file_btn.clicked.connect(self.choise_progect_file)
        self.delite_btn.clicked.connect(self.delite_prod_from_listwidget)
        self.end_applic_btn.clicked.connect(self.add_info_to_word)
        self.standardsizes()
        self.module_widget.hide()
        self.choise_file_widget.hide()
        # self.fill_test_data()  # Заполнение тестовыми данными
        self.load_requester_name()
        self.create_completer()

    def load_requester_name(self) -> None:
        requester_dir = "requester_info"
        requester_file_path = os.path.join(requester_dir, "requester_name.txt")
        if os.path.exists(requester_file_path):
            with open(requester_file_path, "r") as file:
                who_make = file.read().strip()
                self.lineEdit_name_of_ro.setText(who_make)

    def standardsizes(self) -> None:
        height_list = ["2700"]
        width_list = ["2440", "2800", "3000"]
        len_list = ["4880", "6000", "8000", "12000"]
        unists_msrm = ["м", "мм"]
        self.units_of_msrm_combobox.addItems(unists_msrm)
        self.len_combobox.addItems(len_list)
        self.width_combobox.addItems(width_list)
        self.height_combobox.addItems(height_list)

    def constr_solution_changed(self) -> None:
        crnt_constr_solution = self.cnstr_solut_combobox.currentText()
        if crnt_constr_solution == "Модуль (транспак)":
            self.module_widget.show()
            self.other_gabarit_groupBox.hide()
            self.is_it_standart_checkbox.setChecked(True)
        else:
            self.other_gabarit_groupBox.show()
            self.module_widget.hide()

    def gabarits_checkbox_changed_state(self) -> None:
        if self.is_it_standart_checkbox.isChecked():
            self.other_gabarit_groupBox.hide()
            self.standart_gabarits_groupBox.show()
        else:
            self.standart_gabarits_groupBox.hide()
            self.other_gabarit_groupBox.show()

    def have_progect(self) -> None:
        self.files_glob = []
        if self.is_drawings_checkbox.isChecked():
            self.choise_file_widget.show()
            self.quatn_files_lbl.clear()
            self.quatn_files_lbl.setText("0")
            self.ok_lbl.hide()
        else:
            self.choise_file_widget.hide()

    def change_product_in_list(self) -> None:
        if self.listWidget_products.currentItem():
            inf_prod = self.listWidget_products.currentItem().text()
            list_info_about_prod = self.dict_product_and_info[inf_prod]
            self.lineEdit_product_name.clear()
            self.lineEdit_product_name.setText(list_info_about_prod[0])
            index = self.cnstr_solut_combobox.findText(list_info_about_prod[1], QtCore.Qt.MatchContains)
            self.cnstr_solut_combobox.setCurrentIndex(index)
            self.quant_sb.setValue(list_info_about_prod[2])
            if list_info_about_prod[3]:
                ind_len = self.len_combobox.findText(list_info_about_prod[4], QtCore.Qt.MatchContains)
                self.len_combobox.setCurrentIndex(ind_len)
                ind_w = self.width_combobox.findText(list_info_about_prod[5], QtCore.Qt.MatchContains)
                self.width_combobox.setCurrentIndex(ind_w)
                ind_h = self.height_combobox.findText(list_info_about_prod[6], QtCore.Qt.MatchContains)
                self.height_combobox.setCurrentIndex(ind_h)
            else:
                self.len_dblspb.setValue(list_info_about_prod[4])
                self.width_dblspb.setValue(list_info_about_prod[5])
                self.height_dblspb.setValue(list_info_about_prod[6])
                self.units_of_msrm_combobox.findText(list_info_about_prod[7], QtCore.Qt.MatchContains)

    def create_completer(self) -> None:
        list_of_standard_prod = ["Вагон-дом", "Вагон КПП", "Вагон ЛНК", "Вагон-медпост", "Вагон-пункт обогрева",
                            "Вагон-санузел", "Вагон-сварочный пост", "Вагон-сушилка", "Вагон-учебный класс",
                            "Вагон-штаб", "Вагон-прачечная", "Вагон-станция", "Вагон-лаборотория", "Вагон-культцентр",
                            "Вагон-мастерская", "Вагон-сауна", "Вагон-душевая", "Вагон-специального назначения",
                            "Рефрижератор 'Заморозка' 40ф", "Рефрижератор 'Заморозка' 20ф", "Вагон-склад (теплый)",
                            "Вагон-склад (холодный)", "Администратино хозяйственный блок (1 этаж, 30 осей)",
                            "Общежитие на 244 человека (2 этажа, 30 осей)", "ПН-ПН6 (нижний пояс)",
                            "ПС1-ПС6 (пояс средний)", "ПВ1-ПВ5 (пояс верхний)", "К1-К3 (колонна)"
                            ]
        completer = QCompleter(list_of_standard_prod, self.lineEdit_product_name)
        completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.lineEdit_product_name.setCompleter(completer)  # Устанавливает QCompleter в поле ввода

    def add_product_to_listwidget(self) -> None:
        product_name = self.lineEdit_product_name.text()
        crnt_constr_solution = self.cnstr_solut_combobox.currentText()
        quant_of_prod = self.quant_sb.value()
        if self.is_it_standart_checkbox.isChecked():
            len_p = self.len_combobox.currentText()
            width_p = self.width_combobox.currentText()
            height_p = self.height_combobox.currentText()
            units_m = "мм"
        else:
            len_p = self.len_dblspb.value()
            width_p = self.width_dblspb.value()
            height_p = self.height_dblspb.value()
            units_m = self.units_of_msrm_combobox.currentText()
        if quant_of_prod > 0 and product_name:
            indf_prod_name = f"{self.product_index} - {product_name}: {quant_of_prod} шт."
            self.dict_product_and_info[indf_prod_name] = [product_name, crnt_constr_solution, quant_of_prod,
                                                          self.is_it_standart_checkbox.isChecked(),
                                                          len_p, width_p, height_p, units_m
                                                          ]
            self.listWidget_products.addItem(indf_prod_name)
            self.product_index += 1
        else:
            if quant_of_prod == 0:
                error = "Не указано количество продукта!"

            else:
                error = "Укажите название изделия"
            self.MainWindow = ErrorAddReport(error)
            self.MainWindow.show()

    def choise_progect_file(self) -> None:
        self.ok_lbl.hide()
        try:
            filename, filetype = QFileDialog.getOpenFileName(self,
                                                             "Выбрать файл",
                                                             ".",
                                                             "PDF Files (*.pdf);;PNG Files (*.png);;"
                                                             "JPG Files (*.jpg)")
            path_to_file = "{}".format(filename)
            self.ok_lbl.show()
            self.files_glob.append(path_to_file)
            quant = int(self.quatn_files_lbl.text()) + 1
            self.quatn_files_lbl.clear()
            self.quatn_files_lbl.setText(str(quant))

        except Exception as e:
            logging.error(f"Ошибка при выборе файлов: {str(e)}")
            self.MainWindow = ErrorAddReport("Ошибка при выборе файлов")
            self.MainWindow.show()

    def delite_prod_from_listwidget(self) -> None:
        try:
            if self.listWidget_products.selectedItems():
                inf_prod = self.listWidget_products.currentItem().text()
                title = "Подтверждение удаления позиции"
                text = f"Вы уверены, что хотите удалить {inf_prod}?"
                msg = MessageDialogWindow(title, text)  # Вызываем текстовое сообщение с подтверждением удаления
                if msg.confirm_message() == 1:
                    self.dict_product_and_info.pop(inf_prod)
                    self.listWidget_products.clear()
                    if self.dict_product_and_info.keys():
                        self.listWidget_products.addItems(self.dict_product_and_info.keys())
            else:
                error = 'Выберите позицию в списке.'
                self.MainWindow = ErrorAddReport(error)
                self.MainWindow.show()
        except Exception as e:
            logging.error(f"Ошибка при удалении изделия: {str(e)}")
            self.MainWindow = ErrorAddReport("Ошибка при удалении изделия")
            self.MainWindow.show()

    def add_info_to_word(self) -> None:
        error = ""
        try:
            object_name = self.lineEdit_object.text()
            special_text = self.textEdit_speical.toPlainText()
            who_make = self.lineEdit_name_of_ro.text()
            has_a_progect = self.is_drawings_checkbox.isChecked()
            quant_of_dox = self.quatn_files_lbl.text()
            if object_name and who_make:
                if self.dict_product_and_info.keys():
                    product_info = [values for key, values in self.dict_product_and_info.items()]
                    data = {"object_name": object_name,
                            "special_text":special_text,
                            "who_make":who_make,
                            "products": product_info,
                            "has_a_progect": has_a_progect,
                            "quant_of_dox": quant_of_dox
                    }
                    requester_dir = "requester_info"
                    if not os.path.exists(requester_dir):
                        os.makedirs(requester_dir)
                    requester_file_path = os.path.join(requester_dir, "requester_name.txt")
                    with open(requester_file_path, "w") as file:
                        file.write(who_make)
                    exporter = ExportToWord(data)
                    exporter.create_directory_structure()
                    save_path = exporter.create_word_document()
                    # Merge Word and PDF with other files
                    # Convert Word document to PDF
                    pdf_output_dir = os.path.dirname(save_path)
                    docx2pdf.convert(save_path, pdf_output_dir)
                    pdf_file = os.path.join(pdf_output_dir, os.path.splitext(os.path.basename(save_path))[0] + '.pdf')
                    if not os.path.exists(pdf_file):
                        error = "Ошибка конвертации Word в PDF"
                    else:
                        # Initialize PDF merger
                        try:
                            merger = PdfMerger()

                            # Add Word document PDF to PDF merger
                            merger.append(pdf_file)
                            # Add other files to PDF merger
                            if self.files_glob:
                                for file_path in self.files_glob:
                                    print(file_path)
                                    file_extension = os.path.splitext(file_path)[1].lower()
                                    if file_extension == '.pdf':
                                        merger.append(file_path)
                                    elif file_extension in ['.png', '.jpg', '.jpeg']:
                                        # Convert image file to PDF and append to merger
                                        img = Image.open(file_path)
                                        img_temp_pdf = os.path.splitext(file_path)[0] + '.pdf'
                                        img.save(img_temp_pdf, 'PDF')
                                        merger.append(img_temp_pdf)
                                        os.remove(img_temp_pdf)  # Remove temporary PDF image file

                            # Save merged PDF file
                            merged_pdf_path = os.path.join(exporter.export_dir,
                                                           f"Заявка в производство {datetime.now().strftime('%Y%m%d')}.pdf")
                            merger.write(merged_pdf_path)
                            merger.close()
                            # Open the merged PDF file
                            if os.name == 'nt':  # Windows
                                os.startfile(merged_pdf_path)
                            elif os.name == 'posix':  # macOS or Linux
                                subprocess.call(('open' if sys.platform == 'darwin' else 'xdg-open', merged_pdf_path))
                        except Exception as e:
                            logging.error(f"Ошибка при сборке PDF файла: {str(e)}")
                            error = "Ошибка при создании pdf"

                else:
                    error = "Ни одно изделие не добавлено в список!"
            else:
                if not object_name:
                    error = "Укажите наименование объекта!"

                else:
                    error = "Не указан заявитель!"
        except Exception as e:
            logging.error(f"Ошибка при добавлении информации в документ Word: {str(e)}")
            error = "Ошибка при создании документа"
        if error:
            self.MainWindow = ErrorAddReport(error)
            self.MainWindow.show()

    def fill_test_data(self) -> None:
        # Filling object name
        self.lineEdit_object.setText("Тестовый объект")

        # Filling product details
        product_details = [
            ("Ангар КСТ 1", "Изделие", 5, False, 100, 18, 5, "м"),
            ("Общежите", "Рамное", 3, False, 7000, 2800, 3000, "мм"),
            ("Транспак", "Модуль (транспак)", 2, True, "8000", "3000", "2700", "мм")
        ]

        for product in product_details:
            product_name = product[0]
            crnt_constr_solution = product[1]
            quant_of_prod = product[2]
            is_standart = product[3]
            len_p = product[4]
            width_p = product[5]
            height_p = product[6]
            units_m = product[7]
            indf_prod_name = f"{self.product_index} - {product_name}: {quant_of_prod} шт."
            self.dict_product_and_info[indf_prod_name] = [product_name, crnt_constr_solution, quant_of_prod,
                                                          is_standart, len_p, width_p, height_p, units_m]
            self.listWidget_products.addItem(indf_prod_name)
            self.product_index += 1

        # Filling special requirements
        self.textEdit_speical.setPlainText("Тестовые особые требования")

        # Filling who make
        self.lineEdit_name_of_ro.setText("Тестовый заявитель")

        # Selecting some files
        # self.files_glob = ["path/to/test/file1.pdf", "path/to/test/file2.png"]

        # Updating labels
        #self.quatn_files_lbl.setText(str(len(self.files_glob)))
        #self.ok_lbl.setText("Файлы выбраны")
        #self.ok_lbl.show()


class ExportToWord:
    def __init__(self, data):
        self.data = data
        self.base_dir = "Заявки в производство"
        self.today_date = datetime.now().strftime("%d.%m.%Y")  # "%Y-%m-%d"
        self.export_dir = None

    def create_directory_structure(self):
        if not os.path.exists(self.base_dir):
            os.makedirs(self.base_dir)

        date_dir = os.path.join(self.base_dir, self.today_date)
        if not os.path.exists(date_dir):
            os.makedirs(date_dir)

        self.export_dir = os.path.join(date_dir, self.generate_unique_id())
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)

    def generate_unique_id(self):
        return f"UID-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    def create_word_document(self):
        template_path = 'templates/default.docx'
        document = Document(template_path)
        # Set font
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        # Add heading
        heading = document.add_heading('Заявка на производство', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add object information
        document.add_heading('Объект:', level=1)
        document.add_paragraph(self.data['object_name'])

        # Add products information
        document.add_heading('Прошу произвести следующие изделия:', level=1)

        # Add table
        table = document.add_table(rows=1, cols=6)
        list_of_headers = ['Название изделия', 'Конструктивное решение', 'Кол-во, шт.', 'Длина', 'Ширина', 'Высота']
        hdr_cells = table.rows[0].cells
        for index in range(len(list_of_headers)):
            hdr_cells[index].text = list_of_headers[index]
            hdr_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tcPr = hdr_cells[index]._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

        for product in self.data['products']:
            product_name = product[0]
            crnt_constr_solution = product[1]
            quant_of_prod = product[2]
            is_standart = product[3]
            len_p = product[4]
            width_p = product[5]
            height_p = product[6]
            units_m = product[7]

            row_cells = table.add_row().cells
            row_cells[0].text = product_name
            row_cells[1].text = crnt_constr_solution
            row_cells[2].text = str(quant_of_prod)
            row_cells[3].text = f'{len_p} {units_m}'
            row_cells[4].text = f'{width_p} {units_m}'
            row_cells[5].text = f'{height_p} {units_m}'
            # Set border for each cell in the row
            for cell in row_cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # Add special requirements
        document.add_heading('Особые требования:', level=1)
        document.add_paragraph(self.data['special_text'])
        if self.data['has_a_progect']:
            # quant_of_dox
            document.add_paragraph(f"К заявке прилагаю эскизы в кол-ве: {self.data['quant_of_dox']}")

        # Add requester information
        document.add_heading('Заявитель:', level=1)
        document.add_paragraph(self.data['who_make'])
        document.add_heading('Дата составления заявки:', level=1)
        document.add_paragraph(self.today_date)

        # Save document
        save_path = os.path.join(self.export_dir, 'Заявка на производство.docx')
        document.save(save_path)
        return save_path


def application():
    app = QtWidgets.QApplication(sys.argv)
    app.setWindowIcon(QIcon('images/applications.png'))
    MainWindow = MainProgramMenu()
    MainWindow.show()
    sys.exit(app.exec_())


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    application()

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
